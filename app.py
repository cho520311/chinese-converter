import streamlit as st
import re
from io import BytesIO
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- 核心邏輯：保持與 convert.py 一致的聲調顏色判定 ---
def get_tone_color(py_text):
    py = py_text.lower().strip()
    # 1. 第五聲優先判定 (藍色)
    if re.search(r'5$', py) or any(c in py for c in ['â', 'ê', 'î', 'ô', 'û', '̂', 'ˆ', '^']):
        return RGBColor(0, 0, 255)
    # 2. 入聲判定 (以 p, t, k 結尾) -> 紅色
    if py.endswith(('p', 't', 'k')):
        return RGBColor(255, 0, 0)
    # 3. 聲調符號與數字標調 (2, 3, 4, 6, 7, 8) -> 紅色
    marks = ['á', 'à', 'ā', 'ǎ', 'í', 'ì', 'ī', 'ǐ', 'ú', 'ù', 'ū', 'ǔ', 'é', 'è', 'ě', 'ó', 'ò', 'ō', 'ǒ', '̍', '́', '̀', '̌', '̄']
    if any(c in py for c in marks) or re.search(r'[234678]$', py):
        return RGBColor(255, 0, 0)
    # 4. 預設判定：第一聲 (藍色)
    return RGBColor(0, 0, 255)

def set_cell_margins_zero(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for m in ['top', 'left', 'bottom', 'right']:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), '100') # 增加內部邊距讓格子變寬
        node.set(qn('w:type'), 'dxa')
        mar.append(node)
    tcPr.append(mar)

def create_row_table(doc, row_data):
    if not row_data: return
    table = doc.add_table(rows=2, cols=len(row_data))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = True # 讓 Word 根據內容加長格子
    
    for row in table.rows:
        row.allow_break_across_pages = False

    for idx, (hanzi, pinyin) in enumerate(row_data):
        c1 = table.cell(0, idx)
        set_cell_margins_zero(c1)
        c1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p1.add_run(pinyin)
        run1.font.size = Pt(11)
        run1.font.name = 'Times New Roman'
        run1.font.color.rgb = get_tone_color(pinyin)
        run1.bold = True

        c2 = table.cell(1, idx)
        set_cell_margins_zero(c2)
        c2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(hanzi)
        run2.font.size = Pt(20)
        run2.font.name = '標楷體'
        run2._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

    spacer = doc.add_paragraph()
    spacer.paragraph_format.line_spacing = Pt(12)

# --- Streamlit 介面美學 ---
st.set_page_config(page_title="漢字音標轉換工具", page_icon="✨", layout="centered")

st.markdown("""
    <style>
    .main-title { font-size: 36px !important; font-weight: 800; color: #1E3A8A; text-align: center; margin-bottom: 20px; }
    .section-header { font-size: 22px !important; font-weight: 600; color: #475569; margin-top: 20px; margin-bottom: 10px; }
    
    /* 溫馨提示：清淡風格 */
    .info-box { background-color: #F8FAFC; padding: 20px; border-radius: 12px; border: 1px solid #E2E8F0; font-size: 18px; color: #334155; }
    
    /* 幾何圖形樣式 */
    .geo-icon { display: inline-block; width: 15px; height: 15px; margin-right: 10px; }
    .blue-square { background-color: #0000FF; border-radius: 2px; }
    .red-circle { background-color: #FF0000; border-radius: 50%; }
    .gold-triangle { width: 0; height: 0; border-left: 8px solid transparent; border-right: 8px solid transparent; border-bottom: 15px solid #F59E0B; background-color: transparent !important; }

    /* 格式範例：字體極大 */
    .example-box { 
        background-color: #FFFFFF; padding: 30px; border-radius: 15px; border: 2px solid #F1F5F9; 
        font-size: 32px !important; font-weight: bold; text-align: center; color: #1E3A8A; line-height: 1.8;
    }
    
    .upload-label { color: #1E3A8A; font-size: 24px !important; font-weight: bold; margin-top: 30px; }

    /* 下載區塊：加長長度並維持清淡 */
    .download-container {
        background-color: #F1F5F9;
        padding: 40px 60px; /* 增加內距讓框框看起來更長 */
        border-radius: 15px;
        text-align: center;
        margin-top: 20px;
        width: 100%;
    }
    
    div.stDownloadButton > button {
        background-color: #1E3A8A !important;
        color: white !important;
        font-size: 20px !important;
        font-weight: bold !important;
        border-radius: 10px !important;
        padding: 15px 0px !important;
        width: 100% !important;
    }
    </style>
    """, unsafe_allow_html=True)

st.markdown('<div class="main-title">✨ 漢字音標轉換工具</div>', unsafe_allow_html=True)

# 💡 格式範例
st.markdown('<div class="section-header">💡 格式範例</div>', unsafe_allow_html=True)
st.markdown('<div class="example-box">為(uî) 樂(lók) 當(tong) 及(kíp) 時(sî)<br>何(hô) 能(nîng) 待(tǎi) 來(lâi) 茲(tsir)</div>', unsafe_allow_html=True)

# 📢 溫馨提示：幾何圖形原色修正
st.markdown('<div class="section-header">📢 溫馨提示</div>', unsafe_allow_html=True)
st.markdown("""
    <div class="info-box">
        <span class="geo-icon blue-square"></span> 系統將自動根據聲調為音標著色（藍色/紅色）。<br>
        <span class="geo-icon red-circle"></span> 轉換後的 Word 檔將維持標楷體排版。<br>
        <span class="geo-icon gold-triangle"></span> 若有純文字行，系統會自動置中呈現。
    </div>
""", unsafe_allow_html=True)

st.divider()

st.markdown('<div class="upload-label">📥 選擇您的 TXT 檔案並上傳</div>', unsafe_allow_html=True)
uploaded_file = st.file_uploader("", type="txt", label_visibility="collapsed")

if uploaded_file is not None:
    stringio = uploaded_file.getvalue().decode("utf-8")
    lines = stringio.splitlines()

    doc = Document()
    doc.styles['Normal'].font.name = '標楷體'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

    for i, line in enumerate(lines):
        matches = re.findall(r'([\u4e00-\u9fff])\(([^)]+)\)', line)
        if matches:
            create_row_table(doc, matches)
        elif line.strip():
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            doc.add_paragraph()

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    st.success("✅ 轉換完成！")
    
    # 加長後的下載區域
    st.markdown('<div class="download-container">', unsafe_allow_html=True)
    st.download_button(
        label="📥 點擊此處下載產出的 Word 檔案",
        data=file_stream,
        file_name="教材產出.docx"
    )
    st.markdown('</div>', unsafe_allow_html=True)
